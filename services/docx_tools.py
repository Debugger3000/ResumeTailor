from pathlib import Path
from docx import Document


def read_resume(file_path: str) -> dict:
    """Read all text content from a .docx file, organized by paragraph."""
    path = Path(file_path)
    if not path.exists():
        return {'error': f'File not found: {file_path}'}

    doc = Document(path)
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            paragraphs.append({
                'index': i,
                'text': para.text,
                'style': para.style.name,
            })

    return {
        'file': str(path),
        'paragraph_count': len(paragraphs),
        'paragraphs': paragraphs,
    }


def replace_paragraph_text(file_path: str, paragraph_index: int, new_text: str) -> dict:
    """Replace a paragraph's text while preserving its formatting (font, style, etc.)."""
    path = Path(file_path)
    if not path.exists():
        return {'error': f'File not found: {file_path}'}

    doc = Document(path)
    if paragraph_index >= len(doc.paragraphs):
        return {'error': f'Index {paragraph_index} out of range'}

    para = doc.paragraphs[paragraph_index]

    # Preserve formatting from the first run, then replace text
    if para.runs:
        # Clear all runs except the first, update the first's text
        first_run = para.runs[0]
        first_run.text = new_text
        for run in para.runs[1:]:
            run.text = ''
    else:
        para.text = new_text

    doc.save(path)
    return {'ok': True, 'index': paragraph_index, 'new_text': new_text}


# ---
# Tool schema: how we describe these to the model
# This is the JSON Schema format Ollama expects

TOOL_SCHEMAS = [
    {
        'type': 'function',
        'function': {
            'name': 'read_resume',
            'description': 'Read all paragraphs from the resume .docx file. Returns text and indices.',
            'parameters': {
                'type': 'object',
                'properties': {
                    'file_path': {
                        'type': 'string',
                        'description': 'Absolute path to the .docx file',
                    },
                },
                'required': ['file_path'],
            },
        },
    },
    {
        'type': 'function',
        'function': {
            'name': 'replace_paragraph_text',
            'description': 'Replace the text of a specific paragraph in the resume, preserving formatting.',
            'parameters': {
                'type': 'object',
                'properties': {
                    'file_path': {
                        'type': 'string',
                        'description': 'Absolute path to the .docx file',
                    },
                    'paragraph_index': {
                        'type': 'integer',
                        'description': 'Zero-based paragraph index from read_resume',
                    },
                    'new_text': {
                        'type': 'string',
                        'description': 'The new text to replace the paragraph with',
                    },
                },
                'required': ['file_path', 'paragraph_index', 'new_text'],
            },
        },
    },
]


# Map of tool name → callable, for dispatch
TOOL_REGISTRY = {
    'read_resume': read_resume,
    'replace_paragraph_text': replace_paragraph_text,
}