import json
import os
from pathlib import Path
from docx import Document
from ollama import AsyncClient
from const.tailor_prompt import TAILOR_PROMPT, EXTRACT_INDEXES_PROMPT, TAILOR_EXTRACT_PARAGRAPHS_INDEXES_SCHEMA, TAILOR_REPLACE_INDEX_SCHEMA
from const.ai_models import ModelType
import time
from services.ai_model_control.ollama_client import ollama_client
from services.ai_model_control.helpers import run_model

# ollama_client = AsyncClient(host=os.getenv('OLLAMA_HOST', 'http://localhost:11434'))
# OLLAMA_MODEL = os.getenv('OLLAMA_MODEL')

async def extract_applicable_paragraphs(
    paragraphs: list[dict], model_type: ModelType
) -> list[dict]:
    """
    Stage 1: Identify which paragraphs contain job titles or tech/skill mentions.
    Returns [{'index': int, 'original_text': str}, ...] for stage 2.
    """
    user_prompt = json.dumps({
        'resume_paragraphs': [
            {'index': p['index'], 'text': p['text']} for p in paragraphs
        ],
    })

    print(f"=== Stage 1: extracting applicable indexes. Input size: {len(user_prompt)} chars ===")
    start = time.time()

    try:
        parsed, elapsed = await run_model(EXTRACT_INDEXES_PROMPT, user_prompt, TAILOR_EXTRACT_PARAGRAPHS_INDEXES_SCHEMA, model_type)
    except json.JSONDecodeError as e:
        print(f"=== run_model call: tailor resume in place JSON DECODE ERROR: {e} ===")
        return {'index': -1, 'text': 'run_model failed at extract_applicable_paragraphs...'} 

    # elapsed = time.time() - start
    print(f"=== Stage 1 returned in {elapsed:.1f}s ===")
    print(f"=== RAW STAGE 1 OUTPUT: {parsed} ===")

    try:
        # parsed = json.loads(response)
        items = parsed.get('applicable_paragraphs', [])
    except json.JSONDecodeError as e:
        print(f"=== STAGE 1 JSON DECODE ERROR: {e} ===")
        items = []

    # Build an index -> text map from the source of truth (not the model's echo)
    paragraph_lookup = {p['index']: p['text'] for p in paragraphs}

    # Keep only indexes that actually exist; rebuild original_text from source
    applicable = [
        {'index': item['index'], 'original_text': paragraph_lookup[item['index']]}
        for item in items
        if item.get('index') in paragraph_lookup
    ]

    print(f"=== Stage 1 selected {len(applicable)} of {len(paragraphs)} paragraphs ===")
    print(f"=== Selected indexes: {sorted(a['index'] for a in applicable)} ===")
    print(f"=== Stage 1 applicable_paragraphs returned: {applicable}===")

    return applicable



def extract_paragraphs(file_path: Path) -> list[dict]:
    """Pull all non-empty paragraphs with their indices."""
    doc = Document(file_path)
    return [
        {'index': i, 'text': p.text, 'style': p.style.name}
        for i, p in enumerate(doc.paragraphs)
        if p.text.strip()
    ]





#     doc.save(output_path)
async def tailor_resume_in_place(
    input_path: Path,
    output_path: Path,
    applicable_paragraphs: list[dict],
    job_description: str,
    model_type: ModelType,
    approved_skills: list[str] | None = None,
    
) -> tuple[int, str]:
    
    # Group input for model
    user_prompt = json.dumps({
        'job_description': job_description,
        'approved_skills': approved_skills or [],
        'applicable_paragraphs': applicable_paragraphs,
    })

    print(f"=== Calling model. Input prompt size: {len(user_prompt)} chars ===")
    print(f"Input prompt: {user_prompt}")
    print(f"Input prompt: {job_description}")
    # start = time.time()
    

    try:
        parsed, elapsed = await run_model(TAILOR_PROMPT, user_prompt, TAILOR_REPLACE_INDEX_SCHEMA, model_type)
    except json.JSONDecodeError as e:
        print(f"=== run_model call: tailor resume in place JSON DECODE ERROR: {e} ===")
        return -1, 0.0   

    # elapsed = time.time() - start
    #raw_content = response['message']['content']
    print(f"=== Model returned in {elapsed:.1f}s. Output size: {len(parsed)} chars ===")
    print("=== RAW OUTPUT ===")
    # print(response[:3000])
    print("=== END RAW OUTPUT ===")

    try:
        # parsed = json.loads(response)
        changes = parsed.get('changes', [])
        summary = parsed.get('summary', '')
    except json.JSONDecodeError as e:
        print(f"=== JSON DECODE ERROR: {e} ===")
        changes = []
        summary = '(model returned invalid JSON)'

    allowed_indexes = {p['index'] for p in applicable_paragraphs}

    # Apply to a copy
    # if indexes in changes list match then change that line / paragraph with new one.
    doc = Document(input_path)
    applied = 0
    for change in changes:
        idx = change.get('index')
        new_text = change.get('new_text', '')
        if idx is None or idx not in allowed_indexes or idx >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[idx]
        if para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ''
        else:
            para.text = new_text
        applied += 1

    doc.save(output_path)
    print(f"tailor_resume_in_place applied {len(changes)} changes")
    print(f"model summary: {summary}")
    return len(changes), summary