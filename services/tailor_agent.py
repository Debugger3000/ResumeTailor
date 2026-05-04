import json
import os
from pathlib import Path
from docx import Document
from ollama import AsyncClient
from const.model_prompt import MODEL_PROMPT, EXTRACT_INDEXES_PROMPT
import time

ollama_client = AsyncClient(host=os.getenv('OLLAMA_HOST', 'http://localhost:11434'))
OLLAMA_MODEL = os.getenv('OLLAMA_MODEL')

JSON_FORMAT = {
        'type': 'object',
        'properties': {
            'changes': {
                'type': 'array',
                'items': {
                    'type': 'object',
                    'properties': {
                        'index': {'type': 'integer'},
                        'new_text': {'type': 'string'},
                    },
                    'required': ['index', 'new_text'],
                    'additionalProperties': False,
                },
            },
            'summary': {'type': 'string'},
        },
        'required': ['changes', 'summary'],
        'additionalProperties': False,
    }


EXTRACT_INDEXES_SCHEMA = {
    'type': 'object',
    'properties': {
        'indexes': {
            'type': 'array',
            'items': {'type': 'integer'},
        },
    },
    'required': ['indexes'],
    'additionalProperties': False,
}


async def extract_applicable_paragraphs(
    paragraphs: list[dict],
) -> list[dict]:
    """
    Stage 1: Ask the model to identify which paragraphs contain job titles
    or technology/skill mentions. Returns the filtered list of paragraphs
    that stage 2 should consider editing.
    """
    user_prompt = json.dumps({
        'resume_paragraphs': [
            {'index': p['index'], 'text': p['text']} for p in paragraphs
        ],
    })

    print(f"=== Stage 1: extracting applicable indexes. Input size: {len(user_prompt)} chars ===")
    start = time.time()

    response = await ollama_client.chat(
        model=OLLAMA_MODEL,
        messages=[
            {'role': 'system', 'content': EXTRACT_INDEXES_PROMPT},
            {'role': 'user', 'content': user_prompt},
        ],
        format=EXTRACT_INDEXES_SCHEMA,
        options={'temperature': 0.0},
    )

    elapsed = time.time() - start
    raw_content = response['message']['content']
    print(f"=== Stage 1 returned in {elapsed:.1f}s ===")
    print(f"=== RAW STAGE 1 OUTPUT: {raw_content} ===")

    try:
        parsed = json.loads(raw_content)
        indexes = parsed.get('indexes', [])
    except json.JSONDecodeError as e:
        print(f"=== STAGE 1 JSON DECODE ERROR: {e} ===")
        indexes = []

    # Sanity check: drop any indexes the model hallucinated
    valid_indexes = {p['index'] for p in paragraphs}
    indexes = [i for i in indexes if i in valid_indexes]

    # Filter the original paragraph list down to only the flagged indexes
    index_set = set(indexes)
    filtered = [p for p in paragraphs if p['index'] in index_set]

    print(f"=== Stage 1 selected {len(filtered)} of {len(paragraphs)} paragraphs ===")
    print(f"=== Selected indexes: {sorted(index_set)} ===")

    return filtered



def extract_paragraphs(file_path: Path) -> list[dict]:
    """Pull all non-empty paragraphs with their indices."""
    doc = Document(file_path)
    return [
        {'index': i, 'text': p.text, 'style': p.style.name}
        for i, p in enumerate(doc.paragraphs)
        if p.text.strip()
    ]


# def extract_paragraphs(file_path: Path) -> list[dict]:
#     """Pull all paragraphs with their indices, marking which are editable."""
#     doc = Document(file_path)
#     paragraphs = []
#     for i, p in enumerate(doc.paragraphs):
#         text = p.text
#         has_border = _has_bottom_border(p)
#         is_empty = not text.strip()
        
#         paragraphs.append({
#             'index': i,
#             'text': text,
#             'style': p.style.name,
#             'is_empty': is_empty,
#             'has_border': has_border,
#             'editable': not is_empty and not has_border,
#         })
#     return paragraphs


# def _has_bottom_border(paragraph) -> bool:
#     """Check if a paragraph has a bottom border (used for horizontal lines)."""
#     pPr = paragraph._p.find(
#         '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr'
#     )
#     if pPr is None:
#         return False
#     pBdr = pPr.find(
#         '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr'
#     )
#     return pBdr is not None





























#     doc.save(output_path)
async def tailor_resume_in_place(
    input_path: Path,
    output_path: Path,
    paragraphs: list[dict],
    applicable_paragraphs: list[dict],
    job_description: str,
    approved_skills: list[str] | None = None,
) -> tuple[int, str]:
    
    # Group input for model
    user_prompt = json.dumps({
        'job_description': job_description,
        'approved_skills': approved_skills or [],
        'paragraphs': paragraphs,
        'applicable_paragraphs': applicable_paragraphs,
    })

    print(f"=== Calling model. Input prompt size: {len(user_prompt)} chars ===")
    print(f"Input prompt: {user_prompt}")
    print(f"Input prompt: {job_description}")
    start = time.time()

    # get running ollama client and give list of paragraphs to change...
    response = await ollama_client.chat(
        model=OLLAMA_MODEL,
        messages=[
            {'role': 'system', 'content': MODEL_PROMPT},
            {'role': 'user', 'content': user_prompt},
        ],
        format=JSON_FORMAT,
        options={'temperature': 0.0},
    )

    elapsed = time.time() - start
    raw_content = response['message']['content']
    print(f"=== Model returned in {elapsed:.1f}s. Output size: {len(raw_content)} chars ===")
    print("=== RAW OUTPUT ===")
    print(raw_content[:3000])
    print("=== END RAW OUTPUT ===")

    try:
        parsed = json.loads(raw_content)
        changes = parsed.get('changes', [])
        summary = parsed.get('summary', '')
    except json.JSONDecodeError as e:
        print(f"=== JSON DECODE ERROR: {e} ===")
        changes = []
        summary = '(model returned invalid JSON)'

    # Apply to a copy
    # if indexes in changes list match then change that line / paragraph with new one.
    doc = Document(input_path)
    for change in changes:
        idx = change.get('index')
        new_text = change.get('new_text', '')
        if idx is None or idx >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[idx]
        if para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ''
        else:
            para.text = new_text

    doc.save(output_path)
    print(f"tailor_resume_in_place applied {len(changes)} changes")
    print(f"model summary: {summary}")
    return len(changes), summary












async def score_fit(resume_paragraphs: list[dict], job_description: str) -> int:
    """Get a 0-100 fit score from the model."""
    resume_text = '\n'.join(p['text'] for p in resume_paragraphs)
    response = await ollama_client.chat(
        model=OLLAMA_MODEL,
        messages=[
            {
                'role': 'system',
                'content': 'Score resume-job fit from 0 to 100. Respond ONLY with the integer.',
            },
            {
                'role': 'user',
                'content': f"JOB:\n{job_description}\n\nRESUME:\n{resume_text}",
            },
        ],
    )
    try:
        return int(''.join(c for c in response['message']['content'] if c.isdigit())[:3])
    except ValueError:
        return 0
